# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ===== 页面设置 =====
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)

# ===== 辅助函数 =====
def set_font(run, name='微软雅黑', size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading1(text):
    p = add_para('', space_before=18, space_after=8)
    run = p.add_run(text)
    set_font(run, size=18, bold=True, color=(30, 100, 200))

def add_heading2(text):
    p = add_para('', space_before=14, space_after=6)
    run = p.add_run('▌ ' + text)
    set_font(run, size=14, bold=True, color=(30, 100, 200))

def add_heading3(text):
    p = add_para('', space_before=10, space_after=4)
    run = p.add_run('◆ ' + text)
    set_font(run, size=12, bold=True, color=(60, 60, 60))

def add_body(text, indent=False):
    p = add_para('', space_before=3, space_after=3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=11, color=(50, 50, 50))

def add_bullet(text, level=1):
    p = add_para('', space_before=3, space_after=3)
    prefix = '  · ' if level == 1 else '    — '
    p.paragraph_format.left_indent = Cm(0.5 * level)
    run = p.add_run(prefix + text)
    set_font(run, size=11, color=(60, 60, 60))

def add_highlight_box(title, text, bg_color=(240, 247, 255)):
    p = add_para('', space_before=6, space_after=6)
    run = p.add_run(f'【{title}】{text}')
    set_font(run, size=11, bold=False, color=(30, 80, 160))

def add_step(num, title, desc):
    p = add_para('', space_before=6, space_after=3)
    run1 = p.add_run(f'STEP {num}  ')
    set_font(run1, size=12, bold=True, color=(255, 255, 255))
    run2 = p.add_run(title)
    set_font(run2, size=12, bold=True, color=(30, 100, 200))
    p2 = add_para('', space_before=2, space_after=6)
    p2.paragraph_format.left_indent = Cm(0.8)
    run3 = p2.add_run(desc)
    set_font(run3, size=11, color=(80, 80, 80))

def add_divider():
    p = add_para('', space_before=8, space_after=8)
    run = p.add_run('─' * 44)
    set_font(run, size=9, color=(200, 200, 200))

# =============================================
# 正文开始
# =============================================

# 封面标题
p_title = add_para('', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=4)
run = p_title.add_run('企业数字化转型规划系统')
set_font(run, size=26, bold=True, color=(20, 80, 180))

p_sub = add_para('', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
run = p_sub.add_run('——从诊断到落地，全流程规划一站搞定')
set_font(run, size=13, italic=True, color=(100, 100, 120))

p_tag = add_para('', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=20)
run = p_tag.add_run('# 数字化转型  # 企业规划  # SaaS工具  # 管理咨询')
set_font(run, size=10, color=(150, 150, 180))

add_divider()

# =============================================
# 引言
# =============================================
add_heading1('写在前面')

add_body('企业数字化转型，这个词大家已经喊了十年，但真正系统性落地的企业屈指可数。')
add_body('为什么？')
add_body('不是不想转，是不知道怎么转——')
add_bullet('现状在哪里？转型差距有多大？')
add_bullet('愿景是什么？目标怎么设？')
add_bullet('从哪里切入？先干什么后干什么？')
add_bullet('投多少钱？怎么排期？风险怎么管？')
add_body('每一个问题都是一座山。大多数企业在这些问题里打转，最终转型沦为"口号工程"。')
add_body('为此，我们开发了这套「企业数字化转型规划系统」，把咨询方法论嵌入工具，让企业管理者和数字化团队能够系统、高效地完成从诊断到落地的全流程规划。')

add_divider()

# =============================================
# 系统概述
# =============================================
add_heading1('系统是什么')

add_body('这是一套面向企业数字化转型的在线规划平台，支持多人协作、在线使用，覆盖三大核心阶段、二十余个功能模块，帮助企业从零系统性输出一套完整的数字化转型规划方案。')

p = add_para('', space_before=10, space_after=4)
run = p.add_run('三大阶段，环环相扣：')
set_font(run, size=12, bold=True, color=(30, 100, 200))

add_bullet('第一阶段：调研诊断 → 搞清楚现状，找准问题')
add_bullet('第二阶段：总体规划 → 设计目标，搭建蓝图')
add_bullet('第三阶段：实施路径 → 制定计划，排出行动步骤')

add_body('三个阶段的输出相互衔接，最终形成一份完整的数字化转型规划报告，可直接用于董事会汇报或对外交流。')

add_divider()

# =============================================
# 第一阶段
# =============================================
add_heading1('第一阶段：调研诊断')

add_body('数字化转型不能"拍脑袋"，要从摸清家底开始。调研诊断阶段提供了五套专业评估工具：')

add_heading2('① 数字化转型成熟度评测')
add_body('基于行业通用成熟度模型，从五个维度对企业现状打分：')
add_bullet('战略与领导力：高层是否有明确的数字化战略意图？')
add_bullet('组织与文化：是否具备推动变革的组织能力？')
add_bullet('业务流程：核心业务流程数字化程度如何？')
add_bullet('数据与技术：数据资产积累与IT基础设施是否到位？')
add_bullet('生态协同：供应链、客户端的数字化协同能力？')
add_body('系统自动汇总评分，生成雷达图与成熟度等级（初始级/发展级/成熟级/领先级），直观呈现企业所处位置。')

add_heading2('② 数据管理成熟度评估')
add_body('专项评估企业的数据治理能力，涵盖数据架构、数据质量、数据安全、数据应用等维度，输出数据能力短板清单，为后续数据架构设计提供依据。')

add_heading2('③ 智能制造评估')
add_body('针对制造业企业，评估生产、质控、设备、供应链等环节的智能化水平，帮助识别智能制造切入点。')

add_heading2('④ 访谈挖掘机')
add_body('内置结构化访谈框架，支持记录和整理各部门访谈内容，系统自动提炼痛点、需求和改进建议，避免信息遗漏。')

add_heading2('⑤ ADTO评分工具（对标分析）')
add_body('导入行业标杆企业数据，系统自动与企业自评数据进行对比，输出差距分析报告，找准努力方向。')

add_heading2('阶段输出：诊断分析报告')
add_body('完成以上评估后，系统自动生成《数字化转型诊断分析报告》，包括：')
add_bullet('企业现状总结')
add_bullet('成熟度评级与各维度得分')
add_bullet('关键问题清单与根因分析')
add_bullet('转型机会与优先领域建议')

add_divider()

# =============================================
# 第二阶段
# =============================================
add_heading1('第二阶段：总体规划')

add_body('在清晰的现状诊断基础上，总体规划阶段帮助企业构建系统性的转型蓝图，覆盖战略到架构的全层次设计。')

add_heading2('① 转型愿景设计')
add_body('引导管理层提炼数字化转型愿景，明确2～5年内的转型定位与核心价值主张。工具提供愿景框架模板，支持团队共创和在线讨论。')

add_heading2('② 战略目标制定')
add_body('基于愿景，分解具体战略目标，支持OKR/BSC等多种目标管理框架，输出量化的转型指标体系，确保目标可衡量、可追踪。')

add_heading2('③ 七大架构设计（企业架构全视图）')

add_body('系统采用TOGAF企业架构方法论，从七个层面设计转型蓝图：')

add_bullet('业务架构设计：梳理核心业务流程，设计目标业务模型，识别业务能力差距')
add_bullet('数据架构设计：规划数据资产体系，设计主数据、数据中台、数据湖方案')
add_bullet('应用架构设计：绘制应用全景地图，规划核心系统建设与整合路径')
add_bullet('技术架构设计：选择技术路线，设计云原生、微服务、中台等技术平台')
add_bullet('安全架构设计：构建数字安全防护体系，覆盖网络、数据、应用安全')
add_bullet('治理架构设计：建立数字化治理机制，包括组织架构、职责分工、流程规范')

add_body('每个架构模块均提供专业设计工具和模板，支持图形化编辑，输出符合业务要求的架构设计文档。')

add_heading2('阶段输出：总体规划报告')
add_body('汇总以上内容，自动生成《数字化转型总体规划报告》，包含愿景、目标、架构全图、建设重点，可直接用于高层汇报。')

add_divider()

# =============================================
# 第三阶段
# =============================================
add_heading1('第三阶段：实施路径')

add_body('有了规划蓝图，还需要一张清晰的行动地图。实施路径阶段将战略意图转化为可执行的项目计划。')

add_heading2('① 项目优先级矩阵')
add_body('对所有待建项目进行价值 × 可行性二维评估，自动生成优先级矩阵，帮助决策层快速判断"先做什么、后做什么"，避免资源浪费。')

add_heading2('② 实施进度安排')
add_body('支持按年度、季度制定分阶段实施计划，提供甘特图视图，清晰展现项目时间线、里程碑节点和依赖关系。')

add_heading2('③ 系统功能模块规划')
add_body('对各核心系统进行功能模块拆解，明确每个系统的建设范围、功能清单和交付标准，为招标或内部立项提供依据。')

add_heading2('④ 系统集成关系图')
add_body('梳理各系统之间的数据流和集成关系，输出系统集成架构图，避免"烟囱式"建设，确保数据互通。')

add_heading2('⑤ 投资费用估算')
add_body('按项目维度估算建设投入，涵盖软件、硬件、实施服务、培训等费用类别，自动汇总三年总投资，支持分阶段预算规划。')

add_heading2('阶段输出：实施路径报告')
add_body('输出《数字化转型实施路径报告》，包含：')
add_bullet('分阶段实施计划（一年、三年、五年）')
add_bullet('项目优先级排序与理由')
add_bullet('里程碑与关键交付物清单')
add_bullet('三年投资计划与ROI预测')
add_bullet('风险清单与管控措施')

add_divider()

# =============================================
# 系统特色
# =============================================
add_heading1('系统亮点')

add_heading2('方法论内嵌，不需要咨询顾问')
add_body('系统将TOGAF、数字化成熟度模型、ADTO等业界主流方法论内置为工具，使用者无需具备深厚咨询背景，按步骤操作即可输出专业成果。')

add_heading2('三阶段闭环，输出即成果')
add_body('从诊断到规划到实施，三个阶段数据互通，上一阶段的输出自动作为下一阶段的输入，全程无需手动整理，最终一键生成完整报告包。')

add_heading2('多人协作，在线使用')
add_body('支持企业团队多人同时在线使用，管理员可创建账号、审批注册、分配权限，适合跨部门协同推进规划工作。')

add_heading2('管理后台，全程可控')
add_body('系统内置管理后台，超级管理员可管理用户账号（新增、审批、重置密码、删除），查看系统使用日志，确保数据安全可控。')

add_divider()

# =============================================
# 使用流程
# =============================================
add_heading1('怎么用？五步走')

add_step('01', '注册并登录系统', '企业管理员注册账号，提交审批后由超级管理员激活；也可直接申请开通演示账号体验功能。')
add_step('02', '启动调研诊断', '进入「调研诊断」模块，填写成熟度评估问卷，上传企业资料，完成访谈记录，生成诊断报告。')
add_step('03', '开展总体规划', '基于诊断结论，进入「总体规划」模块，依次完成愿景、战略目标、七大架构设计，生成规划报告。')
add_step('04', '制定实施路径', '进入「实施路径」模块，对规划项目进行优先级排序，制定分阶段计划，估算投资，生成实施报告。')
add_step('05', '导出完整方案', '系统整合三阶段成果，一键导出完整版数字化转型规划方案，支持Word/PPT/Excel多种格式。')

add_divider()

# =============================================
# 适用场景
# =============================================
add_heading1('适合谁用')

add_bullet('制造业、零售业、金融业等传统企业的数字化部门负责人')
add_bullet('企业CIO/CDO/信息化总监，需要向董事会呈报数字化规划')
add_bullet('管理咨询顾问，需要高效交付数字化转型咨询报告')
add_bullet('政府和国有企业，推进数字化转型顶层设计')
add_bullet('高校、研究机构，开展数字化转型课题研究')

add_divider()

# =============================================
# 结尾
# =============================================
add_heading1('写在最后')

add_body('数字化转型不是一次性项目，而是一场持续演进的旅程。这套系统的价值，不仅在于帮你输出一份规划报告，更在于帮助团队建立共同的语言体系和方法论框架，让每一个参与者都能在同一张地图上前行。')

add_body('企业数字化转型，需要的不是更多的会议和PPT，而是一套真正能落地的方法和工具。')

p_end = add_para('', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=8)
run = p_end.add_run('—— 欢迎扫码申请演示账号，开启你的数字化规划之旅 ——')
set_font(run, size=12, italic=True, color=(100, 130, 200))

add_divider()

p_foot = add_para('', align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4, space_after=4)
run = p_foot.add_run('本文由企业数字化转型规划系统团队出品  |  转载请注明来源')
set_font(run, size=9, color=(180, 180, 180))

# ===== 保存 =====
output_path = r'C:\Users\adylee\WorkBuddy\20260406101012\digital-transform-system\docs\数字化转型规划系统介绍（公众号文章）.docx'
doc.save(output_path)
print('DONE:', output_path)
